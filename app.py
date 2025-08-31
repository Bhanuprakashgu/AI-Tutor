from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
import os
import json
import tempfile
import base64
import io
from datetime import datetime
import requests
import speech_recognition as sr
import pyttsx3
import threading
import queue
import time
import uuid
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document as DocxDocument
import sqlite3

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.docstore.document import Document
from langchain.chains import RetrievalQA
from langchain.llms.base import LLM
from typing import Optional, List, Mapping, Any

app = Flask(__name__)
CORS(app)

# Configuration
GEMINI_API_KEY = 'AIzaSyB4ohkes95vVHseMy9wej18vaQMTIhY5lU'
CHROMA_PERSIST_DIR = "./chroma_db"
UPLOAD_FOLDER = "./uploads"
DATABASE_FILE = "./ai_tutor.db"

# Create necessary directories
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CHROMA_PERSIST_DIR, exist_ok=True)

# Global variables for TTS
tts_engine = None
tts_queue = queue.Queue()
tts_thread = None

# Global variables for tracking
study_sessions = {}
conversation_count = 0
quiz_count = 0

class GeminiLLM(LLM):
    """Custom LangChain LLM wrapper for Gemini API"""
    
    @property
    def _llm_type(self) -> str:
        return "gemini"
    
    def _call(self, prompt: str, stop: Optional[List[str]] = None) -> str:
        """Call Gemini API with enhanced prompting for comprehensive responses"""
        url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
        headers = {
            'Content-Type': 'application/json',
            'X-goog-api-key': GEMINI_API_KEY
        }
        
        # Enhanced prompt for more comprehensive responses
        enhanced_prompt = f"""
You are a friendly, highly accurate AI tutor and assistant. Speak naturally, be helpful, and handle both casual chat and deep questions.

Tone & style
- Warm, human, and concise by default. Use simple words.
- For casual greetings (“hi”, “hlo”, “hey”), reply in 1–2 short sentences (optional single emoji).
- Avoid filler, flattery, and purple prose.
- Never invent facts; if unsure, say what you know and how to verify.
- Do NOT reveal your hidden chain-of-thought; give a brief reasoning summary only.

Conversation continuity
- If the user continues an old conversation, remember context and respond consistently.
- If the user switches to a new topic, reset naturally and answer that new request directly.
- Always adapt based on whether the question relates to the ongoing thread or starts fresh.

Answering strategy
1) Start with a crisp, direct answer in 1–2 lines.
2) Then give a clear explanation. If simple → one short paragraph + an example.
3) If complex → break into clean sections with short bullet points.
4) Add practical examples/applications when useful.
5) If info is missing → state reasonable assumptions and proceed.
6) If the request is unsafe/illegal → briefly refuse and suggest a safer alternative.

Document handling
- If the user uploads a document, carefully analyze its content.
- Answer based only on the document if the question is document-specific.
- If the document plus external context are both relevant, combine them for a full answer.

Formatting
- Use short paragraphs and tight bullet lists.
- For math: present the final result first, then a tidy explanation.
- For code: minimal, runnable snippets with brief comments.
- Keep lists compact (usually ≤5 bullets) unless asked for more.

Conversation handling
- Small talk: be friendly and brief.
- If the user asks for more depth or examples, expand.
- If the user says “explain like I’m new,” simplify further and add an intuitive example.

Now respond to the user’s message below.

User message: {prompt}
"""

        
        payload = {
            "contents": [
                {
                    "parts": [
                        {
                            "text": enhanced_prompt
                        }
                    ]
                }
            ],
            "generationConfig": {
                "temperature": 0.7,
                "topK": 40,
                "topP": 0.95,
                "maxOutputTokens": 2048
            }
        }
        
        try:
            response = requests.post(url, headers=headers, json=payload)
            response.raise_for_status()
            result = response.json()
            
            if 'candidates' in result and len(result['candidates']) > 0:
                if 'content' in result['candidates'][0] and 'parts' in result['candidates'][0]['content']:
                    return result['candidates'][0]['content']['parts'][0]['text']
        except Exception as e:
            print(f"Gemini API error: {e}")
            return "I'm having trouble connecting to my knowledge base right now. Please try again in a moment."
        
        return "No response generated."
    
    @property
    def _identifying_params(self) -> Mapping[str, Any]:
        return {"api_key": "gemini"}

# Database initialization
def init_database():
    """Initialize SQLite database for tracking"""
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    
    # Create tables
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS documents (
            id TEXT PRIMARY KEY,
            filename TEXT NOT NULL,
            original_name TEXT NOT NULL,
            upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            file_size INTEGER,
            content_preview TEXT
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS study_sessions (
            id TEXT PRIMARY KEY,
            topic TEXT NOT NULL,
            start_time TIMESTAMP,
            end_time TIMESTAMP,
            duration INTEGER,
            date DATE DEFAULT CURRENT_DATE
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS quizzes (
            id TEXT PRIMARY KEY,
            topic TEXT NOT NULL,
            num_questions INTEGER,
            questions TEXT,
            created_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS conversations (
            id TEXT PRIMARY KEY,
            query TEXT NOT NULL,
            response TEXT NOT NULL,
            timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Add quiz scores table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS quiz_scores (
            id TEXT PRIMARY KEY,
            quiz_id TEXT NOT NULL,
            score INTEGER NOT NULL,
            total_questions INTEGER NOT NULL,
            percentage REAL NOT NULL,
            completed_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (quiz_id) REFERENCES quizzes (id)
        )
    ''')
    
    conn.commit()
    conn.close()

# Document processing functions
def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        print(f"Error extracting DOCX text: {e}")
        return ""

def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print(f"Error extracting TXT text: {e}")
        return ""

def process_uploaded_document(file_path, filename):
    """Process uploaded document and add to vector store"""
    file_ext = filename.lower().split('.')[-1]
    
    if file_ext == 'pdf':
        text = extract_text_from_pdf(file_path)
    elif file_ext == 'docx':
        text = extract_text_from_docx(file_path)
    elif file_ext == 'txt':
        text = extract_text_from_txt(file_path)
    else:
        return None, "Unsupported file format"
    
    if not text.strip():
        return None, "No text content found in document"
    
    # Create document object
    doc = Document(
        page_content=text,
        metadata={
            "source": filename,
            "type": "uploaded_document",
            "upload_date": datetime.now().isoformat()
        }
    )
    
    # Split document
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )
    split_docs = text_splitter.split_documents([doc])
    
    # Add to vector store
    try:
        vectorstore.add_documents(split_docs)
        return len(split_docs), None
    except Exception as e:
        return None, f"Error adding to vector store: {str(e)}"

# Initialize RAG components
def initialize_rag():
    """Initialize the RAG pipeline with actual vector database"""
    
    # Sample educational documents
    documents = [
    # Mathematics
    Document(page_content="Algebra is a branch of mathematics dealing with symbols and the rules for manipulating those symbols. Variables represent quantities without fixed values. In algebra, we use letters like x, y, and z to represent unknown numbers. For example, in the equation 2x + 3 = 7, we solve for x by subtracting 3 from both sides to get 2x = 4, then dividing by 2 to find x = 2. Algebraic expressions can be simplified using the distributive property, combining like terms, and following the order of operations (PEMDAS).", metadata={"subject": "mathematics", "topic": "algebra"}),
    Document(page_content="Geometry studies shapes, sizes, relative positions of figures, and properties of space. Euclidean geometry focuses on flat surfaces and straight lines. Key concepts include points (which have no dimension), lines (which extend infinitely in both directions), and planes (flat surfaces that extend infinitely). Angles are measured in degrees, with a full circle being 360 degrees. Common shapes include triangles (3 sides), quadrilaterals (4 sides), pentagons (5 sides), and circles. The Pythagorean theorem states that in a right triangle, a² + b² = c², where c is the hypotenuse.", metadata={"subject": "mathematics", "topic": "geometry"}),
    Document(page_content="Calculus studies continuous change, with differential calculus focused on rates of change and integral calculus on accumulation and areas. Differential calculus deals with derivatives, which represent the rate of change of a function at any given point. For example, if you're driving and your position changes over time, the derivative gives you your velocity. Integral calculus is essentially the reverse process - if you know the velocity, you can find the total distance traveled. The fundamental theorem of calculus connects these two branches, showing that differentiation and integration are inverse operations.", metadata={"subject": "mathematics", "topic": "calculus"}),
    Document(page_content="Linear algebra studies linear equations, matrices, and vector spaces, essential for computer graphics, engineering, and physics. Vectors are quantities that have both magnitude and direction, like velocity or force. Matrices are rectangular arrays of numbers that can represent transformations in space. Matrix multiplication allows us to combine transformations, such as rotating and then scaling an object. Eigenvalues and eigenvectors are special properties of matrices that remain unchanged (except for scaling) when the matrix transformation is applied.", metadata={"subject": "mathematics", "topic": "linear_algebra"}),
    Document(page_content="Probability quantifies the likelihood of events occurring, ranging from 0 (impossible) to 1 (certain). Basic probability is calculated as the number of favorable outcomes divided by the total number of possible outcomes. For example, when rolling a fair six-sided die, the probability of rolling a 3 is 1/6. Independent events don't affect each other (like multiple coin flips), while dependent events do affect each other (like drawing cards without replacement). The addition rule helps calculate the probability of either event A or event B occurring.", metadata={"subject": "mathematics", "topic": "probability"}),
    Document(page_content="Statistics is the study of data: collecting, analyzing, interpreting, and presenting it. Key concepts include mean (average), median (middle value), mode (most frequent value), and standard deviation (measure of spread). The mean is calculated by adding all values and dividing by the number of values. The median is found by arranging values in order and finding the middle value. Standard deviation tells us how spread out the data is from the mean. A normal distribution (bell curve) is a common pattern where most data clusters around the mean.", metadata={"subject": "mathematics", "topic": "statistics"}),

    # Science - Physics, Chemistry, Biology, Earth
    Document(page_content="Physics studies matter, motion, energy, and force. Newton's laws describe how forces affect motion. Newton's First Law states that objects at rest stay at rest, and objects in motion stay in motion, unless acted upon by an external force (inertia). Newton's Second Law gives us F = ma, meaning force equals mass times acceleration. Newton's Third Law states that for every action, there is an equal and opposite reaction. These laws explain everything from why we feel pushed back in an accelerating car to how rockets work in space.", metadata={"subject": "science", "topic": "physics"}),
    Document(page_content="Chemistry studies elements and compounds, their properties, reactions, and the periodic table. Elements are pure substances made of only one type of atom, like hydrogen, oxygen, or gold. The periodic table organizes elements by their atomic number (number of protons). Compounds are formed when different elements bond together, like water (H₂O) which combines hydrogen and oxygen. Chemical reactions involve breaking and forming bonds between atoms. The law of conservation of mass states that matter cannot be created or destroyed in chemical reactions, only rearranged.", metadata={"subject": "science", "topic": "chemistry"}),
    Document(page_content="Biology studies living organisms, their structure, processes, and evolution. Cell theory states all life is made of cells, which are the basic units of life. Prokaryotic cells (like bacteria) lack a nucleus, while eukaryotic cells (like plant and animal cells) have a nucleus containing DNA. Photosynthesis is the process by which plants convert sunlight, carbon dioxide, and water into glucose and oxygen. Evolution by natural selection explains how species change over time, with favorable traits becoming more common in populations.", metadata={"subject": "science", "topic": "biology"}),
    Document(page_content="The ozone layer is a region in Earth's stratosphere that contains high concentrations of ozone (O₃) molecules. This layer acts as Earth's natural sunscreen, absorbing most of the Sun's harmful ultraviolet (UV) radiation before it reaches the surface. Without the ozone layer, life on Earth would be severely damaged by UV radiation, which can cause skin cancer, cataracts, and harm to plants and marine ecosystems. Human-made chemicals like chlorofluorocarbons (CFCs) have depleted the ozone layer, leading to international agreements like the Montreal Protocol to phase out these harmful substances.", metadata={"subject": "science", "topic": "ozone"}),
    Document(page_content="Ecology studies interactions between organisms and their environment, including ecosystems and biodiversity. An ecosystem includes all living organisms (biotic factors) and non-living components (abiotic factors) in a specific area. Food chains show the flow of energy from producers (plants) to primary consumers (herbivores) to secondary consumers (carnivores). Biodiversity refers to the variety of life in an ecosystem - greater biodiversity generally means a more stable and resilient ecosystem. Human activities like deforestation and pollution can disrupt these delicate ecological balances.", metadata={"subject": "science", "topic": "ecology"}),
    Document(page_content="Meteorology is the study of weather, climate patterns, and atmospheric phenomena. Weather refers to short-term atmospheric conditions, while climate describes long-term patterns. The water cycle drives much of our weather: evaporation from oceans and lakes, condensation into clouds, and precipitation as rain or snow. Air pressure differences create wind, and the rotation of Earth causes the Coriolis effect, which influences wind patterns and storm systems. Understanding meteorology helps us predict weather and prepare for severe events like hurricanes and tornadoes.", metadata={"subject": "science", "topic": "meteorology"}),
    Document(page_content="Water (H₂O) is essential for life, acting as a solvent and maintaining biological processes. Water's unique properties include its ability to dissolve many substances (making it the 'universal solvent'), its high specific heat capacity (helping regulate temperature), and its expansion when frozen (which is unusual for most substances). In living organisms, water transports nutrients, removes waste, helps regulate body temperature, and provides the medium for most biochemical reactions. The human body is approximately 60% water, highlighting its critical importance for survival.", metadata={"subject": "science", "topic": "water"}),
    Document(page_content="The human brain controls thought, memory, emotions, motor skills, and sensory perception. The brain has several key regions: the cerebrum (responsible for conscious thought and voluntary actions), the cerebellum (coordinates movement and balance), and the brainstem (controls vital functions like breathing and heart rate). Neurons are the brain's communication cells, transmitting electrical and chemical signals. The brain's plasticity allows it to adapt and form new connections throughout life, which is the basis for learning and memory formation.", metadata={"subject": "science", "topic": "human_brain"}),
    Document(page_content="Climate change refers to long-term shifts in temperature and weather patterns, largely driven by human activity since the Industrial Revolution. The primary cause is increased greenhouse gas emissions, particularly carbon dioxide from burning fossil fuels. These gases trap heat in Earth's atmosphere, leading to global warming. Effects include rising sea levels, more frequent extreme weather events, changes in precipitation patterns, and shifts in ecosystems. Mitigation strategies include transitioning to renewable energy, improving energy efficiency, and protecting forests that absorb carbon dioxide.", metadata={"subject": "science", "topic": "climate_change"}),

    # History
    Document(page_content="Ancient civilizations like Mesopotamia, Egypt, and Greece laid the foundations of modern society. Mesopotamia, located between the Tigris and Euphrates rivers, is often called the 'cradle of civilization' where writing, the wheel, and the first cities developed around 3500 BCE. Ancient Egypt, unified around 3100 BCE, gave us hieroglyphics, monumental architecture like the pyramids, and advances in medicine and mathematics. Ancient Greece contributed democracy, philosophy (with thinkers like Socrates, Plato, and Aristotle), theater, and the Olympic Games. These civilizations established concepts of law, government, art, and science that continue to influence us today.", metadata={"subject": "history", "topic": "ancient_civilizations"}),
    Document(page_content="The Renaissance (14th-17th century) emphasized humanism, art, science, and exploration in Europe. This period marked a 'rebirth' of classical learning and culture after the Middle Ages. Key figures included Leonardo da Vinci (artist, inventor, scientist), Michelangelo (sculptor, painter), and Galileo Galilei (astronomer, physicist). The Renaissance saw the development of perspective in art, the printing press (which spread knowledge), and the scientific method. Humanism placed emphasis on human potential and achievement rather than purely religious concerns. This period also launched the Age of Exploration, with voyages by Columbus, Magellan, and others.", metadata={"subject": "history", "topic": "renaissance"}),
    Document(page_content="World Wars I and II reshaped global politics, economies, and technology, leading to organizations like the UN. World War I (1914-1918) was triggered by the assassination of Archduke Franz Ferdinand and involved complex alliance systems. It introduced new technologies like machine guns, poison gas, and aircraft, resulting in unprecedented casualties. World War II (1939-1945) began with Germany's invasion of Poland and included the Holocaust, the atomic bomb, and major battles across Europe, Africa, and the Pacific. These wars led to the decline of European colonial empires, the rise of the United States and Soviet Union as superpowers, and the creation of international organizations to prevent future conflicts.", metadata={"subject": "history", "topic": "world_wars"}),
    Document(page_content="The Industrial Revolution (18th-19th century) transitioned societies from manual production to mechanized industry. Beginning in Britain around 1760, it was driven by innovations like the steam engine, spinning jenny, and power loom. Factories replaced home-based production, leading to urbanization as people moved from rural areas to cities for work. The revolution improved transportation with canals, railways, and steamships. While it increased productivity and living standards for many, it also created new social problems including poor working conditions, child labor, and environmental pollution. This period laid the foundation for modern industrial society.", metadata={"subject": "history", "topic": "industrial_revolution"}),
    Document(page_content="The Cold War (1947-1991) was a period of geopolitical tension between the US and USSR, marked by nuclear arms race and space competition. Despite being called a 'war,' it involved no direct military conflict between the superpowers. Instead, it featured proxy wars (like in Korea and Vietnam), espionage, propaganda, and competition for global influence. The nuclear arms race created the doctrine of Mutually Assured Destruction (MAD). The Space Race began with the Soviet launch of Sputnik in 1957 and culminated with the US moon landing in 1969. The Cold War ended with the fall of the Berlin Wall in 1989 and the dissolution of the Soviet Union in 1991.", metadata={"subject": "history", "topic": "cold_war"}),

    # Computer Science
    Document(page_content="Python is a high-level, interpreted programming language, widely used in web development, data science, and AI. Created by Guido van Rossum in 1991, Python emphasizes code readability with its clean syntax. It's an interpreted language, meaning code is executed line by line rather than compiled first. Python supports multiple programming paradigms including procedural, object-oriented, and functional programming. Its extensive standard library and third-party packages make it versatile for tasks ranging from web development (Django, Flask) to data analysis (pandas, NumPy) to machine learning (scikit-learn, TensorFlow). Python's simplicity makes it an excellent first programming language.", metadata={"subject": "computer_science", "topic": "python"}),
    Document(page_content="Machine learning is a subset of AI where computers learn patterns from data to make predictions without being explicitly programmed for each task. There are three main types: supervised learning (learning from labeled examples), unsupervised learning (finding patterns in unlabeled data), and reinforcement learning (learning through trial and error with rewards). Common algorithms include linear regression (predicting continuous values), decision trees (making decisions based on features), and neural networks (inspired by brain structure). Machine learning powers recommendation systems, image recognition, natural language processing, and autonomous vehicles.", metadata={"subject": "computer_science", "topic": "machine_learning"}),
    Document(page_content="Algorithms are step-by-step procedures for solving problems, fundamental in computer science. They must be precise, unambiguous, and finite (eventually terminate). Algorithm efficiency is measured by time complexity (how execution time grows with input size) and space complexity (how memory usage grows). Big O notation describes worst-case performance: O(1) is constant time, O(n) is linear time, O(n²) is quadratic time. Common algorithms include sorting (bubble sort, merge sort, quicksort), searching (linear search, binary search), and graph traversal (breadth-first search, depth-first search). Good algorithm design can make the difference between a program that runs in seconds versus hours.", metadata={"subject": "computer_science", "topic": "algorithms"}),
    Document(page_content="Data structures organize data for efficient access and manipulation, including arrays, lists, trees, and graphs. Arrays store elements in contiguous memory locations, allowing fast access by index but fixed size. Linked lists use pointers to connect elements, allowing dynamic size but slower access. Stacks follow Last-In-First-Out (LIFO) principle, useful for function calls and undo operations. Queues follow First-In-First-Out (FIFO) principle, useful for scheduling tasks. Trees organize data hierarchically, with binary search trees enabling fast searching. Hash tables provide average O(1) access time using hash functions to map keys to array indices.", metadata={"subject": "computer_science", "topic": "data_structures"}),
    Document(page_content="Artificial intelligence (AI) simulates human intelligence, including learning, reasoning, and problem-solving. AI can be narrow (designed for specific tasks like chess or image recognition) or general (human-level intelligence across all domains, which doesn't exist yet). Machine learning is a subset of AI that learns from data. Deep learning uses neural networks with many layers to process complex patterns. Natural language processing enables computers to understand and generate human language. Computer vision allows machines to interpret visual information. AI applications include virtual assistants, autonomous vehicles, medical diagnosis, and game playing.", metadata={"subject": "computer_science", "topic": "artificial_intelligence"}),

    # Technology
    Document(page_content="The Internet is a global network enabling communication, information sharing, and e-commerce. It began as ARPANET in the 1960s, designed to maintain communication during potential nuclear attacks. The World Wide Web, created by Tim Berners-Lee in 1989, made the Internet user-friendly with web browsers and hyperlinks. The Internet uses packet switching to break data into small pieces that travel independently and reassemble at the destination. Key protocols include TCP/IP (transmission control), HTTP (web pages), and DNS (domain name system). The Internet has revolutionized commerce, education, entertainment, and social interaction globally.", metadata={"subject": "technology", "topic": "internet"}),
    Document(page_content="Blockchain is a decentralized ledger technology used for cryptocurrencies and secure transactions. It consists of blocks of data linked together using cryptographic hashes, creating an immutable chain. Each block contains a timestamp, transaction data, and a hash of the previous block. The decentralized nature means no single authority controls the blockchain - instead, a network of computers (nodes) maintains copies. Consensus mechanisms like Proof of Work ensure agreement on valid transactions. Beyond cryptocurrencies like Bitcoin, blockchain applications include supply chain tracking, digital identity verification, smart contracts, and voting systems.", metadata={"subject": "technology", "topic": "blockchain"}),
    Document(page_content="Robotics combines engineering, AI, and computer science to build autonomous machines. Robots typically have sensors (to perceive their environment), actuators (to move and manipulate objects), and control systems (to process information and make decisions). Industrial robots perform repetitive tasks in manufacturing with high precision. Service robots assist humans in tasks like cleaning, delivery, or healthcare. Humanoid robots are designed to resemble and interact with humans. Key challenges include navigation in complex environments, object manipulation, and human-robot interaction. Advances in AI, particularly machine learning, are making robots more adaptable and intelligent.", metadata={"subject": "technology", "topic": "robotics"}),
    Document(page_content="Quantum computing leverages quantum mechanics principles to process information in ways classical computers cannot. While classical computers use bits (0 or 1), quantum computers use quantum bits (qubits) that can exist in superposition (both 0 and 1 simultaneously). Quantum entanglement allows qubits to be correlated in ways that classical systems cannot achieve. This enables quantum computers to potentially solve certain problems exponentially faster than classical computers, particularly in cryptography, optimization, and simulation of quantum systems. However, quantum computers are extremely sensitive to environmental interference and require near absolute zero temperatures to operate.", metadata={"subject": "technology", "topic": "quantum_computing"}),
    Document(page_content="Cybersecurity is the practice of protecting systems, networks, and data from digital attacks. Common threats include malware (viruses, worms, trojans), phishing (fraudulent emails to steal information), and denial-of-service attacks (overwhelming systems with traffic). Defense strategies include firewalls (filtering network traffic), encryption (scrambling data), antivirus software, and regular security updates. The CIA triad represents core security principles: Confidentiality (keeping information secret), Integrity (ensuring data accuracy), and Availability (ensuring systems remain accessible). As our dependence on digital systems grows, cybersecurity becomes increasingly critical for individuals, businesses, and governments.", metadata={"subject": "technology", "topic": "cybersecurity"}),

    # Literature & Arts
    Document(page_content="William Shakespeare (1564-1616) was an English playwright and poet, widely regarded as the greatest writer in the English language. His works include 37 plays and 154 sonnets that explore universal themes of love, power, jealousy, betrayal, and human nature. Famous tragedies include Hamlet (a prince's quest for revenge), Macbeth (ambition's destructive power), and Romeo and Juliet (star-crossed lovers). His comedies like A Midsummer Night's Dream and Much Ado About Nothing feature wit, mistaken identities, and happy endings. Shakespeare invented over 1,700 words and phrases still used today, and his works continue to be performed and adapted worldwide.", metadata={"subject": "literature", "topic": "shakespeare"}),
    Document(page_content="Poetry is a form of literary expression using rhythm, rhyme, and imagery to convey emotions and ideas. Key elements include meter (rhythmic pattern), rhyme scheme (pattern of rhyming words), and literary devices like metaphor (comparing unlike things), simile (comparison using 'like' or 'as'), and alliteration (repetition of initial sounds). Different forms include sonnets (14-line poems), haikus (three-line Japanese poems), and free verse (no regular pattern). Poetry can tell stories (narrative), express personal feelings (lyric), or describe scenes (descriptive). Great poets like Emily Dickinson, Robert Frost, and Maya Angelou have used poetry to explore themes of nature, identity, love, and social justice.", metadata={"subject": "literature", "topic": "poetry"}),
    Document(page_content="The Impressionist movement in painting (1860s-1880s) emphasized light, color, and everyday scenes, featuring artists like Claude Monet and Pierre-Auguste Renoir. Impressionists painted outdoors (en plein air) to capture natural light and its changing qualities throughout the day. They used loose brushstrokes and pure colors placed side by side, allowing the viewer's eye to blend them. This was revolutionary compared to the detailed, studio-based academic painting of the time. Monet's Water Lilies series and Renoir's Luncheon of the Boating Party exemplify the movement's focus on light, atmosphere, and modern life. Impressionism influenced all subsequent modern art movements.", metadata={"subject": "arts", "topic": "impressionism"}),

    # Geography
    Document(page_content="Mount Everest, standing at 29,032 feet (8,849 meters), is the highest mountain on Earth, located in the Himalayas on the border between Nepal and China (Tibet). Known as Sagarmatha in Nepali and Chomolungma in Tibetan, it was first successfully climbed by Edmund Hillary and Tenzing Norgay in 1953. The mountain continues to grow about 4 millimeters per year due to tectonic plate movement. Climbing Everest is extremely dangerous due to altitude sickness, avalanches, extreme weather, and the 'death zone' above 26,000 feet where oxygen levels are insufficient to sustain human life for extended periods. Despite the risks, hundreds of climbers attempt the summit each year.", metadata={"subject": "geography", "topic": "mount_everest"}),
    Document(page_content="The Amazon Rainforest, covering about 2.1 million square miles across nine South American countries (primarily Brazil), is the largest tropical rainforest and is known for its incredible biodiversity and role in global climate regulation. Often called the 'lungs of the Earth,' it produces about 20% of the world's oxygen and absorbs massive amounts of carbon dioxide. The Amazon is home to an estimated 10% of all known species, including jaguars, pink river dolphins, and over 1,300 bird species. Indigenous peoples have lived in the Amazon for thousands of years. Unfortunately, deforestation for agriculture, logging, and development threatens this critical ecosystem.", metadata={"subject": "geography", "topic": "amazon_rainforest"}),
    Document(page_content="The Sahara Desert, covering about 3.6 million square miles, is the largest hot desert in the world, located in North Africa. It spans across multiple countries including Morocco, Algeria, Tunisia, Libya, Egypt, Mauritania, Mali, Niger, Chad, and Sudan. Despite its harsh conditions, the Sahara supports various life forms including camels, fennec foxes, and drought-resistant plants. The desert features diverse landscapes including sand dunes (ergs), rocky plateaus (hamadas), and gravel plains (regs). Historically, the Sahara was crossed by trade routes connecting sub-Saharan Africa with the Mediterranean, facilitating the exchange of gold, salt, and other goods.", metadata={"subject": "geography", "topic": "sahara_desert"}),

    # General Knowledge & Misc
    Document(page_content="The solar system consists of the Sun and all celestial objects bound to it by gravity, including eight planets, their moons, asteroids, comets, and cosmic dust. The Sun, a medium-sized star, contains 99.86% of the system's mass and provides the energy that sustains life on Earth. The planets, in order from the Sun, are Mercury, Venus, Earth, Mars (terrestrial planets with solid surfaces), Jupiter, Saturn, Uranus, and Neptune (gas giants). The asteroid belt between Mars and Jupiter contains thousands of rocky objects. Comets, often called 'dirty snowballs,' originate from the outer solar system and develop tails when approaching the Sun.", metadata={"subject": "general_knowledge", "topic": "solar_system"}),
    Document(page_content="Gravity is the fundamental force that attracts objects toward each other, giving weight to physical objects and governing planetary motion. Newton's law of universal gravitation states that every particle attracts every other particle with a force proportional to their masses and inversely proportional to the square of the distance between them. Einstein's theory of general relativity describes gravity not as a force, but as the curvature of spacetime caused by mass and energy. Gravity keeps planets in orbit around the Sun, causes tides on Earth due to the Moon's gravitational pull, and determines the structure of the universe on the largest scales.", metadata={"subject": "general_knowledge", "topic": "gravity"}),
    Document(page_content="Nutrition is the process by which living organisms obtain and use food to support growth, repair, and energy production. Essential nutrients include carbohydrates (primary energy source), proteins (building blocks for tissues), fats (energy storage and cell membranes), vitamins (organic compounds needed in small amounts), minerals (inorganic substances like calcium and iron), and water. A balanced diet includes fruits and vegetables (vitamins, minerals, fiber), whole grains (complex carbohydrates), lean proteins (meat, fish, beans), and healthy fats (nuts, olive oil). Poor nutrition can lead to malnutrition, obesity, diabetes, heart disease, and other health problems.", metadata={"subject": "general_knowledge", "topic": "nutrition"}),
    Document(page_content="Vaccination is a method to stimulate immunity against diseases by introducing a harmless form of the pathogen (weakened, killed, or parts of the organism). When vaccinated, the immune system recognizes the antigen and produces antibodies and activates immune cells. If later exposed to the actual disease, the immune system can quickly recognize and fight it. Vaccines have eliminated smallpox globally and nearly eliminated polio. Common vaccines protect against measles, mumps, rubella, influenza, hepatitis, and COVID-19. Herd immunity occurs when enough people in a community are vaccinated, protecting those who cannot be vaccinated due to medical conditions.", metadata={"subject": "general_knowledge", "topic": "vaccination"}),
]

    
    # Split documents
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )
    
    split_docs = text_splitter.split_documents(documents)
    
    # Create embeddings
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )
    
    # Create vector store
    vectorstore = Chroma.from_documents(
        documents=split_docs,
        embedding=embeddings,
        persist_directory=CHROMA_PERSIST_DIR
    )
    
    # Create LLM
    llm = GeminiLLM()
    
    # Create retrieval chain
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),
        return_source_documents=True
    )
    
    return qa_chain, vectorstore

# Initialize TTS engine
def initialize_tts():
    """Initialize text-to-speech engine"""
    global tts_engine, tts_thread
    
    try:
        tts_engine = pyttsx3.init()
        # Configure TTS settings
        voices = tts_engine.getProperty('voices')
        if voices:
            # Try to use a female voice if available
            for voice in voices:
                if 'female' in voice.name.lower() or 'zira' in voice.name.lower():
                    tts_engine.setProperty('voice', voice.id)
                    break
        
        tts_engine.setProperty('rate', 150)  # Speed
        tts_engine.setProperty('volume', 0.8)  # Volume
        
        # Start TTS worker thread
        tts_thread = threading.Thread(target=tts_worker, daemon=True)
        tts_thread.start()
        
        print("✅ TTS engine initialized successfully")
        return True
    except Exception as e:
        print(f"❌ TTS initialization failed: {e}")
        return False

def tts_worker():
    """Worker thread for TTS processing"""
    global tts_engine
    
    while True:
        try:
            text = tts_queue.get()
            if text is None:  # Shutdown signal
                break
            
            if tts_engine:
                tts_engine.say(text)
                tts_engine.runAndWait()
                
        except Exception as e:
            print(f"TTS error: {e}")
        finally:
            tts_queue.task_done()

def speak_text(text):
    """Add text to TTS queue"""
    try:
        if tts_engine:
            tts_queue.put(text)
    except Exception as e:
        print(f"Error adding text to TTS queue: {e}")

# Initialize speech recognition
def initialize_speech_recognition():
    """Initialize speech recognition"""
    try:
        recognizer = sr.Recognizer()
        microphone = sr.Microphone()
        
        # Adjust for ambient noise
        with microphone as source:
            recognizer.adjust_for_ambient_noise(source)
        
        print("✅ Speech recognition initialized successfully")
        return recognizer, microphone
    except Exception as e:
        print(f"❌ Speech recognition initialization failed: {e}")
        return None, None

# Enhanced quiz generation function
def generate_quiz_questions(topic, num_questions):
    """Generate quiz questions using Gemini API with enhanced prompting"""
    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
    headers = {
        'Content-Type': 'application/json',
        'X-goog-api-key': GEMINI_API_KEY
    }
    
    prompt = f"""
    Create a comprehensive educational quiz about {topic} with exactly {num_questions} multiple-choice questions.
    
    Requirements:
    - Each question should test understanding, not just memorization
    - Include a mix of difficulty levels (easy, medium, hard)
    - Provide 4 answer options (A, B, C, D) for each question
    - Make sure only one answer is clearly correct
    - Cover different aspects of the topic
    - Questions should be educational and informative
    
    Format your response as a JSON array with this exact structure:
    [
        {{
            "question": "Question text here?",
            "options": ["A) Option 1", "B) Option 2", "C) Option 3", "D) Option 4"],
            "correct": "A) Correct option text",
            "explanation": "Brief explanation of why this answer is correct"
        }}
    ]
    
    Topic: {topic}
    Number of questions: {num_questions}
    """
    
    payload = {
        "contents": [
            {
                "parts": [
                    {
                        "text": prompt
                    }
                ]
            }
        ],
        "generationConfig": {
            "temperature": 0.7,
            "topK": 40,
            "topP": 0.95,
            "maxOutputTokens": 4096
        }
    }
    
    try:
        response = requests.post(url, headers=headers, json=payload)
        response.raise_for_status()
        result = response.json()
        
        if 'candidates' in result and len(result['candidates']) > 0:
            content = result['candidates'][0]['content']['parts'][0]['text']
            
            # Extract JSON from the response
            import re
            json_match = re.search(r'\[.*\]', content, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                questions = json.loads(json_str)
                return questions
            else:
                # Fallback: try to parse the entire content as JSON
                questions = json.loads(content)
                return questions
                
    except Exception as e:
        print(f"Quiz generation error: {e}")
        return None

# Initialize components
print("🚀 Initializing AI Tutor components...")
init_database()
qa_chain, vectorstore = initialize_rag()
speech_recognizer, microphone = initialize_speech_recognition()
tts_initialized = initialize_tts()

# Routes
@app.route('/')
def index():
    """Serve the main page"""
    return render_template('index.html')

@app.route('/api/chat', methods=['POST'])
def chat():
    """Handle chat requests"""
    global conversation_count
    
    try:
        data = request.get_json()
        query = data.get('message', '').strip()
        
        if not query:
            return jsonify({'error': 'No message provided'}), 400
        
        # Get response from RAG chain
        result = qa_chain({"query": query})
        response = result['result']
        
        # Store conversation in database
        conversation_count += 1
        conn = sqlite3.connect(DATABASE_FILE)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO conversations (id, query, response) VALUES (?, ?, ?)",
            (str(uuid.uuid4()), query, response)
        )
        conn.commit()
        conn.close()
        
        return jsonify({
            'response': response,
            'sources': [doc.metadata.get('source', 'Unknown') for doc in result.get('source_documents', [])]
        })
        
    except Exception as e:
        print(f"Chat error: {e}")
        return jsonify({'error': 'Failed to process message'}), 500

@app.route('/api/speech/recognize', methods=['POST'])
def recognize_speech():
    """Handle speech recognition"""
    try:
        if not speech_recognizer or not microphone:
            return jsonify({'error': 'Speech recognition not available'}), 500
        
        # Get audio data from request
        audio_data = request.files.get('audio')
        if not audio_data:
            return jsonify({'error': 'No audio data provided'}), 400
        
        # Save audio to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_file:
            audio_data.save(temp_file.name)
            
            # Recognize speech
            with sr.AudioFile(temp_file.name) as source:
                audio = speech_recognizer.record(source)
                text = speech_recognizer.recognize_google(audio)
                
            # Clean up
            os.unlink(temp_file.name)
            
            return jsonify({'text': text})
            
    except sr.UnknownValueError:
        return jsonify({'error': 'Could not understand audio'}), 400
    except sr.RequestError as e:
        return jsonify({'error': f'Speech recognition service error: {e}'}), 500
    except Exception as e:
        print(f"Speech recognition error: {e}")
        return jsonify({'error': 'Speech recognition failed'}), 500

@app.route('/api/speech/synthesize', methods=['POST'])
def synthesize_speech():
    """Handle text-to-speech synthesis"""
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        if not tts_initialized:
            return jsonify({'error': 'Text-to-speech not available'}), 500
        
        # Add text to TTS queue
        speak_text(text)
        
        return jsonify({'success': True})
        
    except Exception as e:
        print(f"TTS error: {e}")
        return jsonify({'error': 'Text-to-speech failed'}), 500

@app.route('/api/documents/upload', methods=['POST'])
def upload_document():
    """Handle document upload"""
    try:
        if 'files' not in request.files:
            return jsonify({'error': 'No files provided'}), 400
        
        files = request.files.getlist('files')
        results = []
        
        for file in files:
            if file.filename == '':
                continue
            
            # Secure filename
            filename = secure_filename(file.filename)
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            
            # Save file
            file.save(file_path)
            
            # Process document
            chunks_added, error = process_uploaded_document(file_path, filename)
            
            if error:
                results.append({'filename': filename, 'error': error})
                continue
            
            # Store in database
            file_size = os.path.getsize(file_path)
            doc_id = str(uuid.uuid4())
            
            conn = sqlite3.connect(DATABASE_FILE)
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO documents (id, filename, original_name, file_size) VALUES (?, ?, ?, ?)",
                (doc_id, filename, file.filename, file_size)
            )
            conn.commit()
            conn.close()
            
            results.append({
                'filename': filename,
                'chunks_added': chunks_added,
                'file_size': file_size
            })
        
        return jsonify({'results': results})
        
    except Exception as e:
        print(f"Document upload error: {e}")
        return jsonify({'error': 'Document upload failed'}), 500

@app.route('/api/documents', methods=['GET'])
def get_documents():
    """Get list of uploaded documents"""
    try:
        conn = sqlite3.connect(DATABASE_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM documents ORDER BY upload_date DESC")
        documents = cursor.fetchall()
        conn.close()
        
        doc_list = []
        for doc in documents:
            doc_list.append({
                'id': doc[0],
                'filename': doc[1],
                'original_name': doc[2],
                'upload_date': doc[3],
                'file_size': doc[4]
            })
        
        return jsonify({'documents': doc_list})
        
    except Exception as e:
        print(f"Get documents error: {e}")
        return jsonify({'error': 'Failed to retrieve documents'}), 500

@app.route('/api/documents/<doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    """Delete a document"""
    try:
        conn = sqlite3.connect(DATABASE_FILE)
        cursor = conn.cursor()
        
        # Get document info
        cursor.execute("SELECT filename FROM documents WHERE id = ?", (doc_id,))
        result = cursor.fetchone()
        
        if not result:
            return jsonify({'error': 'Document not found'}), 404
        
        filename = result[0]
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        
        # Delete file if exists
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # Delete from database
        cursor.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
        conn.commit()
        conn.close()
        
        return jsonify({'success': True})
        
    except Exception as e:
        print(f"Delete document error: {e}")
        return jsonify({'error': 'Failed to delete document'}), 500

@app.route('/api/quiz/generate', methods=['POST'])
def generate_quiz():
    """Generate a quiz"""
    global quiz_count
    
    try:
        data = request.get_json()
        topic = data.get('topic', '').strip()
        num_questions = int(data.get('num_questions', 5))
        
        if not topic:
            return jsonify({'error': 'No topic provided'}), 400
        
        if num_questions < 1 or num_questions > 20:
            return jsonify({'error': 'Number of questions must be between 1 and 20'}), 400
        
        # Generate questions
        questions = generate_quiz_questions(topic, num_questions)
        
        if not questions:
            return jsonify({'error': 'Failed to generate quiz questions'}), 500
        
        # Store quiz in database
        quiz_count += 1
        quiz_id = str(uuid.uuid4())
        
        conn = sqlite3.connect(DATABASE_FILE)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO quizzes (id, topic, num_questions, questions) VALUES (?, ?, ?, ?)",
            (quiz_id, topic, num_questions, json.dumps(questions))
        )
        conn.commit()
        conn.close()
        
        return jsonify({
            'quiz_id': quiz_id,
            'topic': topic,
            'questions': questions
        })
        
    except Exception as e:
        print(f"Quiz generation error: {e}")
        return jsonify({'error': 'Failed to generate quiz'}), 500

@app.route('/api/quiz/submit', methods=['POST'])
def submit_quiz():
    """Submit quiz answers and calculate score"""
    try:
        data = request.get_json()
        quiz_id = data.get('quiz_id')
        answers = data.get('answers', {})
        
        if not quiz_id:
            return jsonify({'error': 'No quiz ID provided'}), 400
        
        # Get quiz from database
        conn = sqlite3.connect(DATABASE_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT questions, topic, num_questions FROM quizzes WHERE id = ?", (quiz_id,))
        result = cursor.fetchone()
        
        if not result:
            return jsonify({'error': 'Quiz not found'}), 404
        
        questions_json, topic, num_questions = result
        questions = json.loads(questions_json)
        
        # Calculate score
        correct_count = 0
        detailed_results = []
        
        for i, question in enumerate(questions):
            question_id = str(i)
            user_answer = answers.get(question_id, '')
            correct_answer = question['correct']
            is_correct = user_answer == correct_answer
            
            if is_correct:
                correct_count += 1
            
            detailed_results.append({
                'question': question['question'],
                'user_answer': user_answer,
                'correct_answer': correct_answer,
                'is_correct': is_correct,
                'explanation': question.get('explanation', '')
            })
        
        percentage = (correct_count / len(questions)) * 100
        
        # Store score in database
        score_id = str(uuid.uuid4())
        cursor.execute(
            "INSERT INTO quiz_scores (id, quiz_id, score, total_questions, percentage) VALUES (?, ?, ?, ?, ?)",
            (score_id, quiz_id, correct_count, len(questions), percentage)
        )
        conn.commit()
        conn.close()
        
        return jsonify({
            'score': correct_count,
            'total': len(questions),
            'percentage': percentage,
            'detailed_results': detailed_results
        })
        
    except Exception as e:
        print(f"Quiz submission error: {e}")
        return jsonify({'error': 'Failed to submit quiz'}), 500

@app.route('/api/quiz/history', methods=['GET'])
def get_quiz_history():
    """Get quiz history"""
    try:
        conn = sqlite3.connect(DATABASE_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quizzes ORDER BY created_date DESC LIMIT 10")
        quizzes = cursor.fetchall()
        conn.close()
        
        quiz_list = []
        for quiz in quizzes:
            quiz_list.append({
                'id': quiz[0],
                'topic': quiz[1],
                'num_questions': quiz[2],
                'created_date': quiz[4]
            })
        
        return jsonify({'quizzes': quiz_list})
        
    except Exception as e:
        print(f"Get quiz history error: {e}")
        return jsonify({'error': 'Failed to retrieve quiz history'}), 500

@app.route('/api/study/start', methods=['POST'])
def start_study_session():
    """Start a study session"""
    try:
        data = request.get_json()
        topic = data.get('topic', '').strip()
        
        if not topic:
            return jsonify({'error': 'No topic provided'}), 400
        
        session_id = str(uuid.uuid4())
        start_time = datetime.now()
        
        # Store in memory for active tracking
        study_sessions[session_id] = {
            'topic': topic,
            'start_time': start_time
        }
        
        # Store in database
        conn = sqlite3.connect(DATABASE_FILE)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO study_sessions (id, topic, start_time) VALUES (?, ?, ?)",
            (session_id, topic, start_time)
        )
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'session_id': session_id,
            'topic': topic,
            'start_time': start_time.isoformat()
        })
        
    except Exception as e:
        print(f"Start study session error: {e}")
        return jsonify({'error': 'Failed to start study session'}), 500

@app.route('/api/study/stop/<session_id>', methods=['POST'])
def stop_study_session(session_id):
    """Stop a study session"""
    try:
        if session_id not in study_sessions:
            return jsonify({'error': 'Study session not found'}), 404
        
        session = study_sessions[session_id]
        end_time = datetime.now()
        duration = int((end_time - session['start_time']).total_seconds())
        
        # Update database
        conn = sqlite3.connect(DATABASE_FILE)
        cursor = conn.cursor()
        cursor.execute(
            "UPDATE study_sessions SET end_time = ?, duration = ? WHERE id = ?",
            (end_time, duration, session_id)
        )
        conn.commit()
        conn.close()
        
        # Remove from active sessions
        del study_sessions[session_id]
        
        # Format duration
        hours = duration // 3600
        minutes = (duration % 3600) // 60
        seconds = duration % 60
        duration_formatted = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
        
        return jsonify({
            'success': True,
            'duration': duration,
            'duration_formatted': duration_formatted
        })
        
    except Exception as e:
        print(f"Stop study session error: {e}")
        return jsonify({'error': 'Failed to stop study session'}), 500

@app.route('/api/study/sessions', methods=['GET'])
def get_study_sessions():
    """Get study session history"""
    try:
        conn = sqlite3.connect(DATABASE_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM study_sessions WHERE end_time IS NOT NULL ORDER BY start_time DESC LIMIT 10")
        sessions = cursor.fetchall()
        conn.close()
        
        session_list = []
        for session in sessions:
            duration = session[4] or 0
            hours = duration // 3600
            minutes = (duration % 3600) // 60
            seconds = duration % 60
            duration_formatted = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
            
            session_list.append({
                'id': session[0],
                'topic': session[1],
                'start_time': session[2],
                'end_time': session[3],
                'duration': duration,
                'duration_formatted': duration_formatted
            })
        
        return jsonify({'sessions': session_list})
        
    except Exception as e:
        print(f"Get study sessions error: {e}")
        return jsonify({'error': 'Failed to retrieve study sessions'}), 500

@app.route('/api/progress', methods=['GET'])
def get_progress():
    """Get learning progress statistics"""
    try:
        conn = sqlite3.connect(DATABASE_FILE)
        cursor = conn.cursor()
        
        # Get counts
        cursor.execute("SELECT COUNT(*) FROM documents")
        documents_count = cursor.fetchone()[0]
        
        cursor.execute("SELECT COUNT(*) FROM quizzes")
        quizzes_count = cursor.fetchone()[0]
        
        cursor.execute("SELECT COUNT(*) FROM study_sessions WHERE end_time IS NOT NULL")
        study_sessions_count = cursor.fetchone()[0]
        
        cursor.execute("SELECT COUNT(*) FROM conversations")
        conversations_count = cursor.fetchone()[0]
        
        # Get total study time
        cursor.execute("SELECT SUM(duration) FROM study_sessions WHERE duration IS NOT NULL")
        total_study_time = cursor.fetchone()[0] or 0
        
        # Format total study time
        hours = total_study_time // 3600
        minutes = (total_study_time % 3600) // 60
        seconds = total_study_time % 60
        total_study_time_formatted = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
        
        # Get unique topics studied
        cursor.execute("SELECT DISTINCT topic FROM study_sessions WHERE end_time IS NOT NULL")
        topics_studied = [row[0] for row in cursor.fetchall()]
        
        # Get recent activity
        recent_activity = []
        
        # Recent documents
        cursor.execute("SELECT original_name, upload_date FROM documents ORDER BY upload_date DESC LIMIT 3")
        for doc in cursor.fetchall():
            recent_activity.append({
                'type': 'document',
                'name': doc[0],
                'date': doc[1]
            })
        
        # Recent quizzes
        cursor.execute("SELECT topic, created_date FROM quizzes ORDER BY created_date DESC LIMIT 3")
        for quiz in cursor.fetchall():
            recent_activity.append({
                'type': 'quiz',
                'name': f"Quiz: {quiz[0]}",
                'date': quiz[1]
            })
        
        # Recent study sessions
        cursor.execute("SELECT topic, start_time FROM study_sessions WHERE end_time IS NOT NULL ORDER BY start_time DESC LIMIT 3")
        for session in cursor.fetchall():
            recent_activity.append({
                'type': 'study',
                'name': f"Study: {session[0]}",
                'date': session[1]
            })
        
        # Sort recent activity by date
        recent_activity.sort(key=lambda x: x['date'], reverse=True)
        recent_activity = recent_activity[:5]  # Keep only top 5
        
        conn.close()
        
        return jsonify({
            'documents_uploaded': documents_count,
            'quizzes_generated': quizzes_count,
            'study_sessions': study_sessions_count,
            'conversations': conversations_count,
            'total_study_time': total_study_time,
            'total_study_time_formatted': total_study_time_formatted,
            'topics_studied': topics_studied,
            'recent_activity': recent_activity
        })
        
    except Exception as e:
        print(f"Get progress error: {e}")
        return jsonify({'error': 'Failed to retrieve progress'}), 500

if __name__ == '__main__':
    print("🎓 AI Tutor System starting...")
    print("📱 Open your browser to: https://127.0.0.1:5000")
    print("🎤 Make sure to allow microphone permissions!")
    app.run(host='0.0.0.0', port=5000, debug=True)

